import re
from io import BytesIO
from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement
from utils.retriver import Retriver
from utils.models import PDFS
from utils.response_generator import generate_response, PROMPT
from sqlalchemy.orm import Session

class TemplateFiller:
    def __init__(self):
        pass

    def extract_angular_context(self, text: str) -> str:
        match = re.search(r"<(.*)>", text, re.DOTALL)
        return match.group(1).strip() if match else ""

    def insert_paragraph_after(self, existing_paragraph, text: str, style=None):
        new_p = OxmlElement("w:p")
        existing_paragraph._element.addnext(new_p)
        para = Paragraph(new_p, existing_paragraph._parent)
        para.text = text
        if style:
            try:
                para.style = style
            except Exception:
                pass
        return para

    def iter_paragraphs(self, doc):
        """Yield all paragraphs including tables."""
        for para in doc.paragraphs:
            yield para
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in self.iter_paragraphs(cell):
                        yield para

    def fill_placeholders(self, doc: Document, placeholders, retrieve_fn):
        paragraphs = list(self.iter_paragraphs(doc))
        i = 0
        while i < len(paragraphs):
            para = paragraphs[i]
            para_text = para.text
            for ph in placeholders:
                if f"<{ph}>" in para_text:
                    context_type = "table" if para._element.getparent().tag.endswith("tc") else "section"
                    content = retrieve_fn(ph, context_type)
                    generated_paragraphs = [p.strip() for p in content.split("\n\n") if p.strip()]
                    if generated_paragraphs:
                        para.text = generated_paragraphs[0]
                        last_para = para
                        for extra_para_text in generated_paragraphs[1:]:
                            last_para = self.insert_paragraph_after(last_para, extra_para_text)
                    break
            i += 1
        return doc

def extract_placeholders(doc: Document):
    """Extract all placeholders from document including tables."""
    placeholders = []

    def iter_paragraphs(doc_or_cell):
        for para in doc_or_cell.paragraphs:
            yield para
        for table in doc_or_cell.tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from iter_paragraphs(cell)

    for para in iter_paragraphs(doc):
        matches = re.findall(r"<(.*?)>", para.text)
        placeholders.extend(matches)

    return placeholders

def retrieve_placeholder_content(
    ph: str,
    context_type: str,
    session: Session,
    user_prompt: str = "",
    process_flow: str = ""
):
    """Retrieve placeholder content using RAG + LLM, with extra user inputs."""
    all_pdfs = session.query(PDFS).all()
    relevant_docs = []

    for pdf in all_pdfs:
        obj = Retriver(document_id=pdf.pdf_uuid)
        relevant_docs += obj.similarity_search(ph)

    relevant_docs.sort(key=lambda x: x.score, reverse=True)

    final_prompt = PROMPT.format(
        user_query=ph,
        relevant_docs=relevant_docs,
        chat_history=[],
        context_type=context_type,
        user_prompt=user_prompt or "",
        flow=process_flow or ""
    )

    print(f"\n[Template Parser Prompt for '{ph}']:\n{final_prompt}\n")

    ans = generate_response(
    docs=relevant_docs,
    query=ph,  # use the placeholder text as query
    chat_history=[],
    context_type=context_type,
    user_prompt=user_prompt,
    flow=process_flow
)

    return ans if isinstance(ans, str) else str(ans)
