"""
Template filling service for processing .docx templates.
"""
import re
from io import BytesIO
from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement
from sqlmodel import Session, select
from utils.models import PDFS
from utils.retriver import Retriver
from services.llm_service import LLMService
from utils.prompt_templates import IMPROVED_PROMPT_TEMPLATE, format_retrieved_chunks
from api.websocket import broadcast_progress_update_sync


class TemplateFiller:
    """Service for filling placeholders in .docx templates."""
    
    def __init__(self, task_id: str = None):
        self.llm_service = LLMService()
        self.task_id = task_id

    def send_call_log(self, service: str, message: str, log_type: str = 'info'):
        """Send a call log via WebSocket."""
        if self.task_id:
            try:
                log_data = {
                    "type": "call_log",
                    "service": service,
                    "message": message,
                    "logType": log_type
                }
                broadcast_progress_update_sync(self.task_id, log_data)
            except Exception as e:
                print(f"Failed to send call log: {e}")

    def extract_angular_context(self, text: str) -> str:
        """Extract content from angular brackets."""
        match = re.search(r"<(.*)>", text, re.DOTALL)
        return match.group(1).strip() if match else ""

    def insert_paragraph_after(self, existing_paragraph, text: str, style=None):
        """Insert a new paragraph after an existing one."""
        new_p = OxmlElement("w:p")
        existing_paragraph._element.addnext(new_p)
        para = Paragraph(new_p, existing_paragraph._parent)
        para.text = text
        if style:
            try:
                para.style = style
            except Exception:
                pass
        return para

    def iter_paragraphs(self, doc):
        """Yield all paragraphs including those in tables."""
        for para in doc.paragraphs:
            yield para
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in self.iter_paragraphs(cell):
                        yield para

    def fill_placeholders(self, doc: Document, placeholders, retrieve_fn):
        """Fill all placeholders in the document."""
        paragraphs = list(self.iter_paragraphs(doc))
        i = 0
        while i < len(paragraphs):
            para = paragraphs[i]
            para_text = para.text
            for ph in placeholders:
                if f"<{ph}>" in para_text:
                    context_type = "table" if para._element.getparent().tag.endswith("tc") else "section"
                    content = retrieve_fn(ph, context_type)
                    generated_paragraphs = [p.strip() for p in content.split("\n\n") if p.strip()]
                    if generated_paragraphs:
                        para.text = generated_paragraphs[0]
                        last_para = para
                        for extra_para_text in generated_paragraphs[1:]:
                            last_para = self.insert_paragraph_after(last_para, extra_para_text)
                    break
            i += 1
        return doc


def extract_placeholders(doc: Document):
    """Extract all placeholders from document including tables."""
    placeholders = []

    def iter_paragraphs(doc_or_cell):
        for para in doc_or_cell.paragraphs:
            yield para
        for table in doc_or_cell.tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from iter_paragraphs(cell)

    for para in iter_paragraphs(doc):
        matches = re.findall(r"<(.*?)>", para.text)
        placeholders.extend(matches)

    return placeholders


def retrieve_placeholder_content(
    ph: str,
    context_type: str,
    session: Session,
    user_prompt: str = "",
    process_flow: str = "",
    task_id: str = None
):
    """Retrieve placeholder content using RAG + LLM with improved prompts."""
    
    def send_call_log(service: str, message: str, log_type: str = 'info'):
        """Send a call log via WebSocket."""
        if task_id:
            try:
                log_data = {
                    "type": "call_log",
                    "service": service,
                    "message": message,
                    "logType": log_type
                }
                broadcast_progress_update_sync(task_id, log_data)
            except Exception as e:
                print(f"Failed to send call log: {e}")
    
    send_call_log("retrieval_service", f"Searching for relevant documents for placeholder: {ph}")
    
    all_pdfs = session.exec(select(PDFS)).all()
    relevant_docs = []

    # Retrieve relevant documents from all PDFs
    for pdf in all_pdfs:
        retriever = Retriver(document_id=pdf.pdf_uuid)
        docs = retriever.similarity_search(ph)
        relevant_docs.extend(docs)

    # Sort by relevance score
    relevant_docs.sort(key=lambda x: x.score, reverse=True)
    
    send_call_log("retrieval_service", f"Found {len(relevant_docs)} relevant documents")

    # Format retrieved chunks for better LLM consumption
    formatted_chunks = format_retrieved_chunks(relevant_docs)

    # Create improved prompt
    send_call_log("llm_service", f"Generating content for placeholder: {ph}")
    
    llm_service = LLMService()
    prompt = IMPROVED_PROMPT_TEMPLATE.format(
        placeholder=ph,
        retrieved=formatted_chunks,
        context_type=context_type,
        user_context=user_prompt or "",
        flow_summary=process_flow or ""
    )

    # Handle image in process flow if present
    image_base64 = None
    if process_flow and process_flow.startswith("data:image/"):
        image_base64 = process_flow.split(",")[1]

    # Get response from LLM
    response = llm_service.query_llm(prompt, image_base64)
    
    send_call_log("llm_service", f"Generated content for placeholder: {ph}")
    
    if response:
        response = str(response).removeprefix("```json").removesuffix('```')
    
    return response if isinstance(response, str) else str(response)
